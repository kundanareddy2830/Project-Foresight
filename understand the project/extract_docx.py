"""
Extract text from all .docx files in the current directory
"""
import os
from docx import Document

def extract_text_from_docx(docx_path):
    """Extract all text from a .docx file"""
    try:
        doc = Document(docx_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading {docx_path}: {str(e)}"

# Get all .docx files in current directory
current_dir = os.path.dirname(os.path.abspath(__file__))
docx_files = [f for f in os.listdir(current_dir) if f.endswith('.docx')]

# Extract text from each file
for docx_file in docx_files:
    docx_path = os.path.join(current_dir, docx_file)
    txt_file = docx_file.replace('.docx', '.txt')
    txt_path = os.path.join(current_dir, txt_file)
    
    print(f"Extracting text from {docx_file}...")
    text = extract_text_from_docx(docx_path)
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(text)
    
    print(f"Saved to {txt_file}")

print("\nExtraction complete!")
